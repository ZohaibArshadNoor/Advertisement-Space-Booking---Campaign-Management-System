import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def create_full_project_guide():
    doc = Document()

    # Page setup - 1 inch margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.header_distance = Inches(0.5)
        section.footer_distance = Inches(0.5)

    # Palette Constants
    COLOR_PRIMARY = RGBColor(15, 23, 42)      # Slate 900
    COLOR_SECONDARY = RGBColor(37, 99, 235)   # Tech Blue (2563EB)
    COLOR_TEAL = RGBColor(13, 148, 136)       # Teal 600
    COLOR_DARK_GRAY = RGBColor(71, 85, 105)   # Slate 600
    COLOR_CHARCOAL = RGBColor(30, 41, 59)     # Slate 800

    HEX_PRIMARY = "1E293B"      # Dark slate for table headers
    HEX_LIGHT_ROW = "F8FAFC"    # Alternate table row
    HEX_CALLOUT_BG = "F1F5F9"   # Callout box background
    HEX_CALLOUT_BORDER = "3B82F6" # Blue border for callout

    # Helper: Set Cell Shading
    def set_cell_shading(cell, hex_color):
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Helper: Set Cell Margins (Padding)
    def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
        tcPr.append(tcMar)

    # Helper: Add Styled Heading 1
    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY
        return p

    # Helper: Add Styled Heading 2
    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = COLOR_SECONDARY
        return p

    # Helper: Add Styled Heading 3
    def add_h3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = COLOR_TEAL
        return p

    # Helper: Add Normal Paragraph
    def add_p(text, bold=False, italic=False, color=COLOR_CHARCOAL, space_after=4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(10.5)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = color
        return p

    # Helper: Add Bullet Item
    def add_bullet(lead_bold, body_text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        run1 = p.add_run(lead_bold)
        run1.font.name = 'Calibri'
        run1.font.size = Pt(10.5)
        run1.font.bold = True
        run1.font.color.rgb = COLOR_PRIMARY

        run2 = p.add_run(body_text)
        run2.font.name = 'Calibri'
        run2.font.size = Pt(10.5)
        run2.font.color.rgb = COLOR_CHARCOAL
        return p

    # Helper: Add Callout Box
    def add_callout(title, body_text):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        cell = table.cell(0, 0)
        cell.width = Inches(6.5)
        set_cell_shading(cell, HEX_CALLOUT_BG)
        set_cell_margins(cell, top=140, bottom=140, left=180, right=180)

        # Left border only (thick blue line)
        tcPr = cell._tc.get_or_add_tcPr()
        borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="36" w:space="0" w:color="{HEX_CALLOUT_BORDER}"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
        tcPr.append(borders)

        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r_title = p.add_run(f"💡 {title}\n")
        r_title.font.name = 'Arial'
        r_title.font.size = Pt(10.5)
        r_title.font.bold = True
        r_title.font.color.rgb = COLOR_SECONDARY

        r_body = p.add_run(body_text)
        r_body.font.name = 'Calibri'
        r_body.font.size = Pt(10)
        r_body.font.color.rgb = COLOR_CHARCOAL

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Helper: Format Table Headers & Rows
    def format_table(table, col_widths, headers, data):
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        # Header Row
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].width = col_widths[i]
            hdr_cells[i].text = h
            set_cell_shading(hdr_cells[i], HEX_PRIMARY)
            set_cell_margins(hdr_cells[i], top=120, bottom=120, left=140, right=140)
            p = hdr_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(9.5)
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

        # Data Rows
        for r_idx, row_data in enumerate(data):
            row_cells = table.add_row().cells
            bg_color = HEX_LIGHT_ROW if r_idx % 2 == 1 else "FFFFFF"
            for c_idx, val in enumerate(row_data):
                row_cells[c_idx].width = col_widths[c_idx]
                row_cells[c_idx].text = str(val)
                set_cell_shading(row_cells[c_idx], bg_color)
                set_cell_margins(row_cells[c_idx], top=90, bottom=90, left=120, right=120)
                p = row_cells[c_idx].paragraphs[0]
                for run in p.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = COLOR_CHARCOAL

        # Set subtle border
        tblPr = table._tbl.tblPr
        borders = parse_xml(f'<w:tblBorders {nsdecls("w")}><w:top w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/><w:insideH w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/><w:insideV w:val="none"/><w:left w:val="none"/><w:right w:val="none"/></w:tblBorders>')
        tblPr.append(borders)

        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Helper: Add Embedded Diagram
    def add_diagram_image(image_filename, caption_text):
        img_path = os.path.join("diagrams", image_filename)
        if os.path.exists(img_path):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run()
            run.add_picture(img_path, width=Inches(6.2))

            cap_p = doc.add_paragraph()
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_p.paragraph_format.space_after = Pt(10)
            cap_run = cap_p.add_run(f"Figure: {caption_text}")
            cap_run.font.name = 'Calibri'
            cap_run.font.size = Pt(9)
            cap_run.font.italic = True
            cap_run.font.color.rgb = COLOR_DARK_GRAY

    # =========================================================================
    # COVER / TITLE PAGE
    # =========================================================================
    p_title_space = doc.add_paragraph()
    p_title_space.paragraph_format.space_before = Pt(36)

    p_badge = doc.add_paragraph()
    p_badge.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_badge = p_badge.add_run("ACADEMIC PROJECT REFERENCE & FACULTY DEFENSE MANUAL")
    r_badge.font.name = 'Arial'
    r_badge.font.size = Pt(9.5)
    r_badge.font.bold = True
    r_badge.font.color.rgb = COLOR_SECONDARY

    p_main_title = doc.add_paragraph()
    p_main_title.paragraph_format.space_before = Pt(6)
    p_main_title.paragraph_format.space_after = Pt(10)
    r_main = p_main_title.add_run("Advertisement Space Booking &\nCampaign Management System")
    r_main.font.name = 'Arial'
    r_main.font.size = Pt(24)
    r_main.font.bold = True
    r_main.font.color.rgb = COLOR_PRIMARY

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(24)
    r_sub = p_sub.add_run("Full Engineering Architecture, File-by-File Breakdown, Initialization Deep-Dive,\n6-Role RBAC Model, Guardrail Mechanics, and Viva Voce Defense Guide")
    r_sub.font.name = 'Calibri'
    r_sub.font.size = Pt(13)
    r_sub.font.color.rgb = COLOR_DARK_GRAY

    # Divider bar
    div_table = doc.add_table(rows=1, cols=1)
    div_cell = div_table.cell(0, 0)
    div_cell.width = Inches(6.5)
    set_cell_shading(div_cell, "2563EB")
    div_cell.paragraphs[0].paragraph_format.space_after = Pt(1)
    doc.add_paragraph().paragraph_format.space_after = Pt(24)

    # Metadata Grid
    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_widths = [Inches(2.2), Inches(4.3)]
    metadata_rows = [
        ("Project Architecture", "Decoupled 3-Tier Enterprise Application (Flask REST API + React 19 Vite SPA)"),
        ("Database Engine", "PostgreSQL 14+ with SQLAlchemy 3.x ORM (22 Relational Tables)"),
        ("Security & Auth", "Flask-JWT-Extended Token Authentication + Granular Role-Based Access (RBAC)"),
        ("Frontend UI Framework", "React 19, React-Bootstrap, React Router v7, Chart.js, Three.js (3D Billboard)"),
        ("Target Evaluation", "Departmental Faculty Presentation, Viva Voce Defense & IT Project Exhibition")
    ]
    for idx, (label, val) in enumerate(metadata_rows):
        c0 = meta_table.rows[idx].cells[0]
        c1 = meta_table.rows[idx].cells[1]
        c0.width, c1.width = meta_widths[0], meta_widths[1]
        c0.text, c1.text = label, val
        set_cell_margins(c0, top=60, bottom=60, left=100, right=100)
        set_cell_margins(c1, top=60, bottom=60, left=100, right=100)
        c0.paragraphs[0].runs[0].font.bold = True
        c0.paragraphs[0].runs[0].font.color.rgb = COLOR_PRIMARY
        c1.paragraphs[0].runs[0].font.color.rgb = COLOR_CHARCOAL
        c0.paragraphs[0].runs[0].font.name = 'Calibri'
        c1.paragraphs[0].runs[0].font.name = 'Calibri'

    doc.add_page_break()

    # Setup Header & Footer on subsequent sections
    header = doc.sections[0].header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hrun = hp.add_run("Advertisement Space Booking System — Faculty Study Guide")
    hrun.font.name = 'Calibri'
    hrun.font.size = Pt(8.5)
    hrun.font.color.rgb = RGBColor(148, 163, 184)

    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    frun = fp.add_run("Confidential — Academic Defense Preparation Manual")
    frun.font.name = 'Calibri'
    frun.font.size = Pt(8.5)
    frun.font.color.rgb = RGBColor(148, 163, 184)

    # =========================================================================
    # SECTION 1: EXECUTIVE OVERVIEW & PROBLEM STATEMENT
    # =========================================================================
    add_h1("1. Executive Overview & Problem Statement")
    
    add_p("The Advertisement Space Booking and Campaign Management System is a commercial-grade digital platform designed to automate and formalize the complete lifecycle of out-of-home (OOH), transit, digital screen, and online advertising inventory management. Historically, the media advertising sector operates on fragmented legacy practices: space availability is checked across manual spreadsheets, rate negotiations occur over unrecorded phone calls, creative assets are transmitted across unstructured email chains, and invoices are manually reconciled with significant operational delay.")
    
    add_p("This project eliminates those operational bottlenecks by delivering an enterprise-ready, double-sided software platform connecting two key stakeholder groups:")
    add_bullet("Advertisers (Clients): ", "Who discover advertising inventory through multi-criteria geographic and dimensional filters, check guaranteed calendar availability, create multi-channel campaigns, request quotations, upload creative assets, make verified payments, and track live performance analytics.")
    add_bullet("Agency Operations Staff: ", "A 5-tier specialized internal team (Sales Executives, Space Managers, Creative Reviewers, Finance Officers, and System Administrators) that oversees inventory health, pricing matrices, artwork validation gates, billing reconciliation, and audit compliance.")

    add_callout(
        "Core Project Objective",
        "To provide a digital, auditable, and mathematically protected workflow where spaces cannot be double-booked, creative assets cannot broadcast without rigorous content and dimensional approval, and campaigns cannot execute without certified financial reconciliation."
    )

    # =========================================================================
    # SECTION 2: END-TO-END WORKFLOW (10-STAGE LIFECYCLE)
    # =========================================================================
    add_h1("2. End-to-End System Workflow (10-Stage Lifecycle)")
    add_p("The system orchestrates a rigorous 10-stage sequential pipeline. Every stage transitions the entities (Campaign, Booking, Creative, Invoice) through strictly validated status state-machines backed by PostgreSQL constraints and service-layer validation.")

    add_diagram_image("workflow_diagram.png", "The 10-Stage Sequential Campaign Lifecycle Pipeline")

    workflow_data = [
        ("1. Registration", "Public / Auth", "Advertiser registers with company name, email, and password. JWT access/refresh tokens issued upon email verification."),
        ("2. Space Search", "Catalog Browser", "Advertiser searches inventory filtered by category (Billboard, Digital Screen, Transit), location (City, GPS), and dimensions."),
        ("3. Availability Check", "Availability Engine", "System executes mathematical interval collision algorithms to verify zero overlapping reservations for the requested dates."),
        ("4. Campaign Creation", "Campaigns Module", "Advertiser creates a campaign container in DRAFT status with defined budget thresholds and strategic objectives."),
        ("5. Quotation & Pricing", "Sales / Quotations", "System applies active Rate Cards, calculating subtotal = base_rate × duration (days). Tax and promotional discounts are computed."),
        ("6. Booking Approval", "Bookings Service", "Booking request moves to PENDING. Space Manager or Sales Executive reviews and approves, reserving the inventory with pessimistic row locking."),
        ("7. Creative Approval", "Creative Quality Gate", "Advertiser uploads media (.png, .mp4, .pdf). Creative Reviewer inspects dimensions, resolution, and compliance before granting APPROVED status."),
        ("8. Invoicing & Payment", "Finance & Billing", "Finance Officer issues official invoice with sales tax. Advertiser pays via Bank Transfer, Card, or Online. Payment reconciliation updates balance due."),
        ("9. Campaign Execution", "Execution Engine", "Once both Creative and Payment gates pass, campaign transitions to ACTIVE status. Space displays start broadcasting."),
        ("10. Completion & Analytics", "Audit & Analytics", "When end_date elapses, status automatically flips to COMPLETED. Real-time impressions, occupancy rates, and financial reports are archived.")
    ]
    format_table(
        doc.add_table(rows=1, cols=3),
        [Inches(1.5), Inches(1.5), Inches(3.5)],
        ["Workflow Stage", "Responsible Module", "Technical & Operational Actions"],
        workflow_data
    )

    # =========================================================================
    # SECTION 3: THE 3 CORE BUSINESS GUARDRAILS & INTEGRITY GATES
    # =========================================================================
    add_h1("3. The 3 Core Business Guardrails & Concurrency Controls")
    add_p("In commercial advertising, software failures lead directly to expensive legal disputes, brand damage, or revenue loss. To guarantee enterprise reliability, the platform enforces 3 Non-Negotiable Business Guardrails at the database and service layers.")

    add_diagram_image("guardrails_diagram.png", "The 3 Non-Negotiable Business Guardrails Enforced by Backend Service Layer")

    add_h2("Guardrail 1: Zero Double-Booking Guarantee (Mathematical Proof & Row Locking)")
    add_p("A single billboard or digital slot cannot be booked by two advertisers for overlapping date ranges under any circumstance. The platform enforces this via two complementary defense layers:")
    add_bullet("1. Interval Intersection Logic: ", "Given an existing booking [S_exist, E_exist] and a requested booking [S_req, E_req], a scheduling collision exists if and only if:\nS_exist <= E_req  AND  E_exist >= S_req\nThe AvailabilityService executes this exact query against the space_availability table. If any record matches, the booking attempt is instantly aborted.")
    add_bullet("2. PostgreSQL Pessimistic Row Locking (.with_for_update()): ", "To eliminate race conditions when two users click 'Book Space' at the exact same millisecond, the backend executes:\nAdvertisingSpace.query.filter_by(id=space_id).with_for_update().first()\nThis places a PostgreSQL row-level exclusive lock on the inventory row. The second transaction is physically queued by PostgreSQL until the first transaction completes its availability check and commits its reservation.")

    add_h2("Guardrail 2: Creative Artwork Approval Gate")
    add_p("An advertiser cannot publish unvetted material to public billboards or transit screens. Even if an advertiser has booked a space and paid in full, the campaign status CANNOT transition to ACTIVE until uploaded media assets receive APPROVED status from a certified Creative Reviewer.")
    add_bullet("Dimensional Validation: ", "Uploaded media is checked against the physical width, height, and aspect ratio of the target advertising space (e.g., 1920x1080 for LED boards).")
    add_bullet("Human Quality Audit: ", "Reviewers verify brand decency, font legibility, copyright standards, and local advertising regulatory compliance. If REJECTED, a mandatory rejection reason is logged, triggering automated notifications for versioned re-upload.")

    add_h2("Guardrail 3: Financial Settlement & Invoice Reconciliation Gate")
    add_p("Space reservation does not grant immediate broadcast rights. The campaign execution engine checks invoice reconciliation before activation:")
    add_bullet("Atomic Balance Tracking: ", "Invoices track subtotal, tax (e.g., 16% sales tax), total_amount, amount_paid, and balance_due. Status transitions from ISSUED -> PARTIALLY_PAID -> PAID.")
    add_bullet("Payment Verification: ", "While online credit card payments reconcile immediately, offline bank transfers and corporate cheques require verification by the Finance Officer role before the invoice balance is cleared.")

    doc.add_page_break()

    # =========================================================================
    # SECTION 4: THE 6 USER ROLES & FUNCTIONALITY MATRIX
    # =========================================================================
    add_h1("4. The 6 User Roles & Comprehensive Functionality Matrix")
    add_p("The platform implements strict Multi-Tenant Role-Based Access Control (RBAC). Every user account belongs to exactly one Role, which defines both database-level permissions (stored as JSON in the roles table) and endpoint-level authorization via Flask decorators.")

    add_diagram_image("roles_matrix.png", "The 6 System Roles & Operational Responsibility Matrix")

    # Detailed breakdown per role
    roles_detail = [
        ("Role 1: Advertiser (External Client)",
         "The commercial client seeking to promote products, services, or events across physical and digital inventory.",
         [
             ("Browse Inventory Catalog", "Filter advertising spaces by Category, Location (City, Address), Base Rate, and Availability dates."),
             ("Check Live Availability", "Interactive calendar view to see free versus occupied slots for any specific space."),
             ("Create & Manage Campaigns", "Create strategic campaign containers (Draft, Pending, Active, Completed) and allocate budget."),
             ("Submit Booking Requests", "Reserve spaces for specific date intervals, receiving instantaneous pricing calculations."),
             ("Upload Creative Assets", "Upload banners, digital video files (.mp4), and print-ready proofs (.pdf) with dimension tags."),
             ("Review Invoices & Make Payments", "View itemized tax invoices, record payment transactions (Bank Transfer / Card), and monitor balances."),
             ("Track Campaign Performance", "Real-time dashboard displaying active bookings, upcoming schedules, and performance metrics.")
         ],
         "Cannot alter space base rates, cannot approve their own creative media, cannot issue official invoices, and cannot view other advertisers' campaigns."),

        ("Role 2: Sales Executive (Agency Commercial)",
         "The internal agency representative responsible for advertiser client onboarding, deal structuring, and quotation management.",
         [
             ("Manage Advertiser Leads", "Create and maintain advertiser company profiles, business contact details, and client portfolios."),
             ("Generate Quotations", "Convert draft campaigns into formal quotations, applying custom corporate discounts and rate card terms."),
             ("Assist Client Bookings", "Initiate space booking requests on behalf of corporate clients who book via telephone or enterprise contract."),
             ("Monitor Sales Pipeline", "Track campaigns across negotiation stages to forecast monthly and quarterly agency billing."),
             ("Resolve Booking Inquiries", "Coordinate with Space Managers to negotiate premium slots during high-demand festival seasons.")
         ],
         "Cannot approve creative artwork, cannot verify bank payment settlements, and cannot modify system user accounts or server settings."),

        ("Role 3: Space Manager (Inventory Operations)",
         "The operations lead responsible for physical inventory health, hardware specifications, and calendar slot management.",
         [
             ("Inventory CRUD Operations", "Add, edit, and retire advertising spaces; specify physical dimensions, resolution, and base daily rates."),
             ("Location & Category Management", "Register geographic locations (Address, City, GPS Latitude/Longitude) and space categories (Billboard, LED Screen, Transit)."),
             ("Rate Card Configuration", "Define tiered pricing models, seasonal peak-rate surcharges, and minimum booking durations."),
             ("Manage Calendar Availability", "Inspect global occupancy schedules; manually block spaces for hardware maintenance, structural repairs, or electrical work."),
             ("Review Pending Bookings", "Accept or decline incoming space booking requests based on operational feasibility.")
         ],
         "Cannot verify financial bank transfers, cannot alter user roles or system permissions, and cannot upload client creative artwork."),

        ("Role 4: Creative Reviewer (Quality Control & Compliance)",
         "The media specialist responsible for vetting all visual and video artwork prior to public display.",
         [
             ("Inspect Creative Queue", "Dedicated dashboard showing all newly uploaded creative assets in PENDING status."),
             ("Verify Technical Specs", "Ensure uploaded media matches physical space aspect ratios, pixel resolutions, and supported formats (.jpg, .png, .mp4)."),
             ("Regulatory & Decency Compliance", "Verify artwork adheres to advertising standards, consumer protection laws, and community decency guidelines."),
             ("Approve / Reject Media", "One-click approval allowing campaign progression, or rejection with required actionable feedback to the advertiser."),
             ("Audit Version Re-Uploads", "Inspect revised artwork submissions to verify that previously flagged defects were corrected.")
         ],
         "Cannot create spaces, cannot book inventory, cannot generate invoices, and cannot modify financial records."),

        ("Role 5: Finance Officer (Accounting & Billing)",
         "The financial officer responsible for commercial billing, tax governance, and payment reconciliation.",
         [
             ("Issue Campaign Invoices", "Generate formal numbered tax invoices derived from confirmed campaign bookings, automatically computing sales tax."),
             ("Record & Verify Payments", "Log customer payments received via bank transfer, cheque, or electronic gateway; cross-reference bank transaction IDs."),
             ("Manage Split / Partial Payments", "Record partial advance installments; track balance_due with automated invoice state reconciliation."),
             ("Monitor Accounts Receivable", "Track aging invoices, flag OVERDUE accounts, and trigger automated payment reminder notifications."),
             ("Export Financial Audit Reports", "Export itemized revenue ledgers, collected sales tax totals, and client settlement summaries.")
         ],
         "Cannot approve or reject creative artwork, cannot modify space physical coordinates, and cannot delete confirmed bookings."),

        ("Role 6: Administrator (Master Governance)",
         "The root system administrator possessing universal oversight across identity, configuration, and security.",
         [
             ("User Account Provisioning", "Create, activate, suspend, or deactivate internal agency staff accounts and external advertiser accounts."),
             ("Dynamic RBAC Configuration", "Inspect and update JSON permission dictionaries assigned to each system role."),
             ("Inspect Immutable Audit Logs", "Review complete forensic audit logs detailing every user login, entity creation, modification, and status change."),
             ("System Health Oversight", "Monitor database connectivity, cache health, storage consumption in uploads/ directory, and API response latencies."),
             ("Emergency Override Authority", "Administrative override capability to resolve operational deadlocks across any module.")
         ],
         "Subject to immutable audit logging; all administrative actions are permanently tracked with user ID, timestamp, and IP address.")
    ]

    for role_title, role_desc, actions, restrictions in roles_detail:
        add_h2(role_title)
        add_p(role_desc, italic=True)
        add_p("Permitted Functional Actions & Operations:", bold=True)
        for act_title, act_desc in actions:
            add_bullet(f"{act_title}: ", act_desc)
        add_p(f"Security Restrictions & Boundaries: {restrictions}", italic=True, color=RGBColor(185, 28, 28), space_after=8)

    # Summary Role Comparison Table
    add_h3("Comparative Capabilities Across All 6 Roles")
    role_summary_headers = ["Functional Domain", "Advertiser", "Sales Exec", "Space Mgr", "Creative Rev", "Finance Off", "Admin"]
    role_summary_rows = [
        ("Inventory Browsing", "Read-Only", "Full Access", "Full CRUD", "Read-Only", "Read-Only", "Full CRUD"),
        ("Space Availability Management", "Check Only", "Check Only", "Full Control", "None", "None", "Full Control"),
        ("Campaign Creation", "Own Only", "Assisted", "None", "None", "None", "Full Control"),
        ("Quotation Preparation", "Accept Only", "Create/Edit", "View Only", "None", "View Only", "Full Control"),
        ("Creative Media Upload", "Upload Own", "Upload", "None", "None", "None", "Full Control"),
        ("Creative Review & Approval", "None", "None", "None", "Full Control", "None", "Full Control"),
        ("Invoice Generation", "View Own", "View Only", "None", "None", "Full Control", "Full Control"),
        ("Payment Verification", "Submit Pay", "View Only", "None", "None", "Full Control", "Full Control"),
        ("User Account Management", "None", "None", "None", "None", "None", "Full Control"),
        ("Audit Logs Inspection", "None", "None", "None", "None", "None", "Full Control")
    ]
    format_table(
        doc.add_table(rows=1, cols=7),
        [Inches(1.8), Inches(0.78), Inches(0.78), Inches(0.78), Inches(0.78), Inches(0.78), Inches(0.8)],
        role_summary_headers,
        role_summary_rows
    )

    doc.add_page_break()

    # =========================================================================
    # SECTION 5: ENTERPRISE SYSTEM ARCHITECTURE
    # =========================================================================
    add_h1("5. Enterprise System Architecture & Technology Stack")
    add_p("The application adopts an enterprise Decoupled Client-Server 3-Tier Architecture. Unlike legacy monolithic web frameworks that tightly couple HTML server templates with backend logic, this architecture completely separates the frontend single-page application (SPA) from the RESTful backend API service.")

    add_diagram_image("system_architecture.png", "Enterprise 3-Tier Client-Server Architecture Overview")

    add_h2("Architectural Layer Breakdown")
    add_bullet("Presentation Layer (Vite + React 19 SPA): ", "Runs entirely inside the client browser. Features 19 domain-isolated component modules, React Router v7 navigation, React-Bootstrap styling, Chart.js KPI data visualizations, and Three.js 3D billboard rendering. Interacts with the backend purely through asynchronous JSON REST calls.")
    add_bullet("API Application Layer (Flask Python Microservices): ", "Stateless REST API organized into 14 modular Flask Blueprints. Implements the Application Factory Pattern, Marshmallow request validation schemas, Flask-JWT-Extended authentication, and centralized error handling middleware.")
    add_bullet("Business Logic & Service Layer: ", "Pure Python service classes (BookingService, AvailabilityService, CampaignService, PaymentService, InvoiceService, MediaService, AuditService) containing all core business logic, validation rules, and transaction boundaries separate from HTTP routing.")
    add_bullet("Persistence Layer (PostgreSQL 14+ & SQLAlchemy ORM): ", "22 strongly-typed relational tables with foreign keys, unique indexes, cascading rules, and timestamp tracking. Managed via Flask-Migrate (Alembic) for automated version-controlled schema migrations.")

    add_h3("Comprehensive Technology Stack Specification")
    tech_headers = ["Layer", "Technology", "Version", "Role & Engineering Responsibility"]
    tech_data = [
        ("Backend Framework", "Python / Flask", "3.14 / 3.x", "Core micro-framework hosting RESTful endpoints & blueprints"),
        ("Database Engine", "PostgreSQL", "14+", "ACID-compliant relational database management system"),
        ("Database ORM", "Flask-SQLAlchemy", "3.x", "Object-Relational Mapping, row-level locking & query builder"),
        ("Schema Migrations", "Flask-Migrate (Alembic)", "4.x", "Version-controlled database schema evolution & upgrade scripts"),
        ("Authentication", "Flask-JWT-Extended", "4.x", "Stateless JSON Web Token (JWT) access and refresh authentication"),
        ("Payload Validation", "Marshmallow", "3.x", "Strict deserialization, schema validation & JSON serialization"),
        ("API Documentation", "Flasgger (Swagger)", "1.0", "Auto-generated interactive OpenAPI specification at /apidocs/"),
        ("Frontend Engine", "React / Vite", "19.2 / 8.2", "Ultra-fast modern Single-Page Application (SPA) development server"),
        ("Frontend Routing", "React Router", "v7.18", "Client-side declarative routing with ProtectedRoute guard rails"),
        ("HTTP Client", "Axios", "1.19", "Promise-based HTTP client with request & response JWT interceptors"),
        ("UI & Component Kit", "Bootstrap / React-Bootstrap", "5.3 / 2.10", "Responsive grid system, accessible modals, cards & form controls"),
        ("Charts & Analytics", "Chart.js / React-Chartjs-2", "4.5 / 5.3", "Interactive bar, line, and doughnut KPI metrics charts"),
        ("3D Graphics Engine", "Three.js / React Three Fiber", "0.185", "Real-time 3D billboard visualizer for advertiser mockup preview")
    ]
    format_table(
        doc.add_table(rows=1, cols=4),
        [Inches(1.3), Inches(1.5), Inches(0.9), Inches(2.8)],
        tech_headers,
        tech_data
    )

    doc.add_page_break()

    # =========================================================================
    # SECTION 6: COMPLETE FILE & FOLDER STRUCTURE EXPLANATION
    # =========================================================================
    add_h1("6. Complete File & Folder Structure Walkthrough")
    add_p("To explain this project to faculty, it is essential to understand the purpose of every folder and key file in the repository. The project is organized cleanly into backend/ and frontend/ directories.")

    add_h2("A. Backend Directory Structure (backend/)")
    add_bullet("backend/run.py: ", "Root execution script for development. Imports create_app() from the app package, instantiates the Flask server, and runs it with debug=True.")
    add_bullet("backend/wsgi.py: ", "Web Server Gateway Interface production entry point. Exposes the raw application callable (app = create_app('production')) for production web servers such as Gunicorn or uWSGI.")
    add_bullet("backend/config.py: ", "Centralized configuration class hierarchy. Contains Base Config, DevelopmentConfig, TestingConfig, and ProductionConfig, reading environment variables (DATABASE_URL, JWT_SECRET_KEY, CORS_ORIGINS) via python-dotenv.")
    add_bullet("backend/create_admin.py: ", "CLI bootstrap utility to seed the initial master Administrator account into PostgreSQL if no admin exists.")
    add_bullet("backend/requirements.txt: ", "Comprehensive Python package manifest pinning exact library versions for reproducible builds.")
    add_bullet("backend/.env & .env.example: ", "Environment configuration files containing local database passwords, JWT signing secrets, and server port allocations.")
    add_bullet("backend/migrations/: ", "Alembic migration repository storing timestamped schema migration scripts (versions/) to synchronize database tables across team machines.")
    add_bullet("backend/uploads/: ", "Local media storage directory where uploaded creative assets (images, videos, PDF proofs) are persisted on disk.")
    add_bullet("backend/tests/: ", "Automated test suite (Pytest) validating business rules, double-booking prevention, role security decorators, and schema validators.")

    add_h2("B. Backend Application Package (backend/app/)")
    add_p("Inside backend/app/, the code is architected into domain modules following Flask's Blueprint pattern:")
    
    backend_folders = [
        ("app/__init__.py", "The Application Factory. Defines create_app(), initializes database extensions, configures Swagger, and registers all 14 domain blueprints."),
        ("app/extensions.py", "Declares unattached Flask extension instances (db = SQLAlchemy(), migrate = Migrate(), jwt = JWTManager()). Prevents circular imports."),
        ("app/models/", "Database ORM layer containing 15 SQLAlchemy model definitions corresponding to the 22 relational database tables."),
        ("app/services/", "Dedicated service layer isolating business rules from HTTP controllers (BookingService, AvailabilityService, PaymentService, etc.)."),
        ("app/common/", "Shared cross-cutting utilities including RBAC decorators (@roles_required), email token services, and date helpers."),
        ("app/errors/", "Centralized error handling middleware capturing AppExceptions, Marshmallow validation errors, and database rollbacks."),
        ("app/seed/", "Database seeding scripts (seed_data.py, seed_demo.py) injecting the 6 default roles and rich demo data for exhibitions."),
        ("app/auth/", "Authentication blueprint: Registration, login, Google OAuth verification, password resets, and JWT token issuance."),
        ("app/spaces/", "Advertising spaces blueprint: Inventory catalog CRUD, categories, geographic locations, and rate card endpoints."),
        ("app/availability/", "Space availability blueprint: Calendar range conflict queries, schedule blocking, and interval validation."),
        ("app/campaigns/", "Campaign management blueprint: Campaign creation, budget allocations, date tracking, and status transitions."),
        ("app/bookings/", "Booking blueprint: Space reservation requests, pessimistic locking conflict checks, and booking approval flows."),
        ("app/creatives/", "Media asset blueprint: Multi-part file uploads to uploads/ directory, asset download streaming, and review approval/rejection."),
        ("app/invoices/", "Billing blueprint: Invoice creation from campaign bookings, sales tax computations, and accounts receivable tracking."),
        ("app/payments/", "Payments blueprint: Payment recording, bank reference logging, and atomic invoice status reconciliation."),
        ("app/dashboard/", "Analytics blueprint: Role-tailored operational KPIs, revenue stats, occupancy percentages, and quick actions."),
        ("app/notifications/", "Notification engine: Event-driven alerts dispatched to users on booking confirmations, approvals, and invoice issuances."),
        ("app/audit_logs/", "Audit blueprint: Read-only access for Administrators to inspect timestamped immutable change records across the platform.")
    ]
    format_table(
        doc.add_table(rows=1, cols=2),
        [Inches(2.0), Inches(4.5)],
        ["Package / Module", "Architectural Role & File Contents"],
        backend_folders
    )

    add_h2("C. Frontend Directory Structure (frontend/)")
    add_bullet("frontend/index.html: ", "Single-Page Application HTML shell. Hosts the <div id='root'></div> mount point where the React virtual DOM attaches.")
    add_bullet("frontend/vite.config.js: ", "Vite build configuration configuring the @vitejs/plugin-react plugin, local development server port, and asset bundling rules.")
    add_bullet("frontend/package.json: ", "Node.js dependency manifest containing React 19, React Router v7, Bootstrap, Axios, Chart.js, and Three.js dependencies.")
    add_bullet("frontend/src/main.jsx: ", "JavaScript entry point. Imports global Bootstrap CSS and executes ReactDOM.createRoot(root).render(<App />).")
    add_bullet("frontend/src/App.jsx: ", "Master routing coordinator. Wraps the application in ThemeProvider and AuthProvider, defining public, protected, and role-restricted routes.")
    add_bullet("frontend/src/context/AuthContext.jsx: ", "Global authentication context provider. Holds the logged-in user object, JWT access token, and role helpers (isAdmin, isAdvertiser).")
    add_bullet("frontend/src/services/apiClient.js: ", "Configured Axios HTTP client with automatic request interceptors injecting JWT Bearer headers and response interceptors handling 401 redirects.")
    add_bullet("frontend/src/layouts/AppShell.jsx: ", "Unified responsive dashboard layout featuring dynamic role-based navigation sidebar, header, theme toggles, and notification bell.")
    add_bullet("frontend/src/features/: ", "Modular feature folders (spaces, campaigns, bookings, creatives, payments, notifications, users, dashboard) each containing dedicated pages and API callers.")

    doc.add_page_break()

    # =========================================================================
    # SECTION 7: CORE ARCHITECTURAL CONCEPTS DEMYSTIFIED (FACULTY FOCUS)
    # =========================================================================
    add_h1("7. Core Architectural Concepts Demystified (Faculty Defense Focus)")
    add_p("Faculty evaluators frequently probe foundational software engineering principles. This section directly answers the core structural questions regarding initiation, schemas, package roots, and architectural patterns.")

    add_diagram_image("initiation_flow.png", "Step-by-Step Application & Database Initiation Execution Flow")

    add_h2("Question 1: What is the Root File used for (run.py vs wsgi.py)?")
    add_p("In Flask applications, the 'Root File' serves as the external bootloader and execution boundary for the Python process:")
    add_bullet("Development Server (run.py): ", "Contains the standard Python idiom:\nif __name__ == '__main__':\n    app = create_app()\n    app.run(debug=True)\nWhen executed directly via `python run.py`, Python sets `__name__` to `'__main__'`. This instantiates the Flask application via the factory and boots the built-in development server with live code-reloading (debug=True).")
    add_bullet("Production WSGI (wsgi.py): ", "In production, Flask's development server is insecure and cannot handle high concurrency. Production environments use a Web Server Gateway Interface (WSGI) server such as Gunicorn, uWSGI, or Nginx. These servers do not execute scripts as `'__main__'`; instead, they import an exposed application variable. `wsgi.py` exposes `app = create_app('production')` directly at the module root.")

    add_h2("Question 2: What is __init__.py used for, and why is it critical?")
    add_p("The `__init__.py` file plays two vital roles in Python and Flask architectures:")
    add_bullet("1. Package Marker: ", "In Python, placing an `__init__.py` file inside a folder informs the Python interpreter that the directory should be treated as an importable package. This allows statements like `from app.models import User` or `from app.auth import auth_bp`.")
    add_bullet("2. Application Factory Pattern (create_app()): ", "In `app/__init__.py`, we encapsulate application creation inside a function `create_app(config_class)`. Why? In naive Flask apps, beginners declare `app = Flask(__name__)` globally at the top of a file. This causes catastrophic circular import errors (e.g., routes import app, app imports models, models import db, db imports app). With the factory pattern, the app is created cleanly on demand, extensions are bound dynamically, blueprints are registered, and isolated test instances can be spawned effortlessly.")
    add_bullet("3. Model Discovery for Migrations: ", "In `app/models/__init__.py`, all 15 model files are imported into a single namespace. When Flask-Migrate (Alembic) runs `flask db migrate`, it inspects SQLAlchemy's `db.metadata`. If models are not imported into `__init__.py`, Alembic will be blind to them and will generate empty migration files!")

    add_h2("Question 3: What are Schemas used for (Marshmallow vs Models)?")
    add_p("A frequent faculty question is: 'Why do you have both models.py and schemas.py? Isn't that redundant?' The answer is that they address completely different architectural concerns:")
    add_bullet("Database Models (SQLAlchemy): ", "Define how data is physically structured and stored inside the PostgreSQL database (table names, columns, data types, primary keys, foreign keys, relationships).")
    add_bullet("Validation Schemas (Marshmallow): ", "Define the contract for incoming and outgoing HTTP JSON payloads across 4 key responsibilities:\n"
               "• Request Validation: Enforces rules BEFORE touching the database (e.g., password length >= 8, valid email format, positive budget decimal).\n"
               "• Deserialization: Safely parses untrusted client JSON strings into native Python objects (e.g., ISO date string '2026-09-01' -> Python date object).\n"
               "• Security (Mass Assignment Prevention): Prevents attackers from injecting unauthorized database fields. If an attacker submits `{\"role\": \"Administrator\"}` on public signup, Marshmallow strips it because it is not declared in `RegistrationSchema`.\n"
               "• Serialization: Converts complex SQLAlchemy database models into clean, sanitized JSON dicts, masking sensitive columns like `password_hash`.")

    add_h2("Question 4: Where and how are we initiating things?")
    add_p("The table below details exactly where, how, and why every major component is initialized across the system:")
    init_headers = ["Component", "Declaration Location", "Binding / Initiation Location", "Initiation Mechanism & Purpose"]
    init_data = [
        ("Flask Application", "app/__init__.py", "Inside create_app() function", "app = Flask(__name__). Instantiated on demand using configuration class."),
        ("SQLAlchemy Database", "app/extensions.py", "app/__init__.py: db.init_app(app)", "db = SQLAlchemy() created unbound, then attached to app. Avoids circular dependencies."),
        ("Database Connection", "config.py (DATABASE_URL)", "On first DB query / transaction", "SQLAlchemy connects lazily via PostgreSQL connection pool over TCP port 5432."),
        ("Alembic Migrations", "app/extensions.py", "app/__init__.py: migrate.init_app(app, db)", "Binds Alembic engine to Flask app and SQLAlchemy database metadata."),
        ("JWT Authentication", "app/extensions.py", "app/__init__.py: jwt.init_app(app)", "Initializes JWTManager to decode and verify incoming Bearer tokens."),
        ("Swagger / OpenAPI", "app/__init__.py", "Swagger(app, config=swagger_config)", "Flasgger parses docstrings and mounts Swagger UI at /apidocs/."),
        ("CORS Security", "app/__init__.py", "CORS(app, origins=..., supports_credentials=True)", "Enables cross-origin AJAX requests from React dev server (http://localhost:5173)."),
        ("Error Handlers", "app/errors/handlers.py", "app/__init__.py: register_error_handlers(app)", "Attaches global error handling middleware for AppExceptions and DB rollbacks."),
        ("Domain Blueprints", "app/<domain>/__init__.py", "app/__init__.py: app.register_blueprint(bp)", "Mounts 14 self-contained API routers (auth, spaces, campaigns, bookings, etc.).")
    ]
    format_table(
        doc.add_table(rows=1, cols=4),
        [Inches(1.2), Inches(1.4), Inches(1.8), Inches(2.1)],
        init_headers,
        init_data
    )

    doc.add_page_break()

    # =========================================================================
    # SECTION 8: FRONTEND CLIENT ARCHITECTURE & FEATURES
    # =========================================================================
    add_h1("8. In-Depth Guide to the Frontend Client Architecture")
    add_p("The frontend is built using React 19 and Vite 8 as a high-performance Single-Page Application (SPA). The application delivers responsive, accessible operational dashboards for all 6 roles while maintaining strict client-side routing guardrails.")

    add_h2("A. Authentication Lifecycle & Token Management (AuthContext.jsx)")
    add_p("User session state is managed globally through React's Context API via AuthContext.jsx:")
    add_bullet("Session Storage: ", "Upon successful login, the backend returns a JWT access token and user metadata object (id, name, email, role). These are saved in browser localStorage.")
    add_bullet("App Initialization: ", "When the user opens the application, useEffect() in AuthProvider checks localStorage for an existing token. If found, it calls /api/auth/me to verify token validity. If expired, it flushes localStorage and resets user state.")
    add_bullet("Role Helpers: ", "Exposes convenience flags like `isAdmin`, `isAdvertiser`, and user role permissions directly to UI components to conditionally render buttons, links, and forms.")

    add_h2("B. Axios Interceptor Architecture (apiClient.js)")
    add_p("All network communication travels through a centralized Axios instance configured with two bidirectional interceptors:")
    add_bullet("1. Request Interceptor: ", "Automatically inspects localStorage before every outgoing HTTP request. If an access_token exists, it attaches the HTTP header:\nAuthorization: Bearer <access_token>\nThis completely eliminates the need for developers to manually pass tokens in component API calls.")
    add_bullet("2. Response Interceptor: ", "Monitors all incoming HTTP responses. If a 401 Unauthorized status is intercepted (indicating token expiration or tampering), it automatically purges localStorage and forces a client redirect to `/login`, safeguarding the user's session.")

    add_h2("C. Protected & Role-Restricted Routing (ProtectedRoute.jsx & AppShell.jsx)")
    add_p("Client-side routes are partitioned into three security levels inside App.jsx:")
    add_bullet("Public Routes: ", "Accessible without authentication: `/` (Landing Page), `/login`, `/register`, `/verify-email`, `/forgot-password`, `/reset-password`.")
    add_bullet("Authenticated Routes: ", "Wrapped in `<ProtectedRoute>`. Verifies `isAuthenticated === true`. If false, redirects to `/login`. Renders the common `<AppShell>` (Sidebar, Header, Profile Menu) and child pages: `/dashboard`, `/spaces`, `/availability`, `/campaigns`, `/bookings`, `/creatives`, `/payments`, `/notifications`.")
    add_bullet("Role-Restricted Routes: ", "Wrapped in `<ProtectedRoute allowedRoles={['Administrator']}>`. If a non-admin attempts to access `/users`, `/audit`, or `/settings`, the component immediately intercepts the navigation and renders `<ForbiddenPage>` (HTTP 403 Equivalent in UI).")

    add_h2("D. Key Operational UI Modules")
    add_bullet("Interactive Availability Calendar: ", "Visual color-coded calendar allowing Space Managers and Advertisers to view booked versus available dates across any inventory item.")
    add_bullet("Creative Upload & Preview Modal: ", "Drag-and-drop file upload with live client image/video thumbnail preview, dimension validation, and submission for review.")
    add_bullet("Real-Time KPI Dashboard: ", "Powered by Chart.js. Renders occupancy percentages, monthly sales revenue curves, campaign stage breakdowns, and pending review counts tailored to each role.")
    add_bullet("Interactive 3D Billboard Visualizer: ", "Engineered with Three.js and React Three Fiber (@react-three/fiber). Allows prospective advertisers to preview how their banner artwork looks when mapped onto a photorealistic 3D billboard structure in an urban setting.")

    doc.add_page_break()

    # =========================================================================
    # SECTION 9: DATABASE SCHEMA & ENTITY RELATIONSHIPS (22 TABLES)
    # =========================================================================
    add_h1("9. Database Schema Reference (22 Relational Tables)")
    add_p("The database architecture consists of 22 strongly-typed relational tables in PostgreSQL. The tables are grouped below by functional domain:")

    db_groups = [
        ("Domain 1: Identity & Access Control", [
            ("roles", "id (PK), name (VARCHAR 50, UNIQUE), permissions (JSON), created_at (TIMESTAMP)"),
            ("users", "id (PK), role_id (FK -> roles.id), advertiser_id (FK -> advertisers.id, NULLABLE), name, email (UNIQUE), password_hash, is_active (BOOL), created_at"),
            ("audit_logs", "id (PK), user_id (FK -> users.id, NULLABLE), action (VARCHAR 50), entity_type, entity_id, old_values (JSON), new_values (JSON), ip_address, created_at")
        ]),
        ("Domain 2: Advertising Inventory & Geo-Location", [
            ("space_categories", "id (PK), name (VARCHAR 100, UNIQUE, e.g. Billboard, Digital Screen, Transit, Banner)"),
            ("locations", "id (PK), name, address, city, latitude (NUMERIC 10,7), longitude (NUMERIC 10,7)"),
            ("advertising_spaces", "id (PK), category_id (FK), location_id (FK), name, description, dimensions, resolution, base_rate (DECIMAL), is_active (BOOL), created_at"),
            ("rate_cards", "id (PK), space_id (FK), name, price (DECIMAL), duration_days (INT), min_booking_days, is_peak (BOOL), created_at"),
            ("space_availability", "id (PK), space_id (FK), booking_id (FK, NULLABLE), start_date (DATE), end_date (DATE), is_booked (BOOL), notes")
        ]),
        ("Domain 3: Advertisers & Client Accounts", [
            ("advertisers", "id (PK), company_name (VARCHAR 150), tax_number, address, website, created_at"),
            ("advertiser_contacts", "id (PK), advertiser_id (FK), name, email, phone, designation, is_primary (BOOL)")
        ]),
        ("Domain 4: Campaigns, Bookings & Reservations", [
            ("campaigns", "id (PK), user_id (FK), advertiser_id (FK), campaign_reference (UNIQUE), name, description, start_date, end_date, budget, status (DRAFT, PENDING, ACTIVE, COMPLETED, CANCELLED)"),
            ("campaign_spaces", "id (PK), campaign_id (FK), space_id (FK), allocated_budget, created_at"),
            ("bookings", "id (PK), campaign_id (FK), space_id (FK), user_id (FK), advertiser_id (FK), booking_reference (UNIQUE), start_date, end_date, total_price, status (PENDING, CONFIRMED, CANCELLED, COMPLETED)")
        ]),
        ("Domain 5: Quotations & Legal Contracts", [
            ("quotations", "id (PK), campaign_id (FK), quotation_number (UNIQUE), subtotal, tax, discount, total_amount, status (DRAFT, SENT, ACCEPTED, REJECTED), valid_until"),
            ("quotation_items", "id (PK), quotation_id (FK), space_id (FK), description, unit_price, quantity, amount"),
            ("contracts", "id (PK), quotation_id (FK), contract_number (UNIQUE), terms_and_conditions, signed_at, status (DRAFT, SIGNED, EXPIRED)")
        ]),
        ("Domain 6: Media Creatives & Content Review", [
            ("creative_assets", "id (PK), campaign_id (FK), uploaded_by (FK -> users.id), file_path, original_filename, file_type, file_size, dimensions, version (INT), status (PENDING, APPROVED, REJECTED)"),
            ("creative_reviews", "id (PK), creative_id (FK), reviewer_id (FK -> users.id), status, comments, reviewed_at")
        ]),
        ("Domain 7: Billing, Invoices & Payments Ledger", [
            ("invoices", "id (PK), campaign_id (FK), advertiser_id (FK), invoice_number (UNIQUE), subtotal, tax, total_amount, amount_paid, balance_due, status (ISSUED, PARTIALLY_PAID, PAID, OVERDUE, CANCELLED), due_date"),
            ("payments", "id (PK), invoice_id (FK), payment_reference (UNIQUE), amount (DECIMAL), payment_method (BANK_TRANSFER, CREDIT_CARD, CHEQUE, ONLINE), transaction_reference, status (COMPLETED, PENDING, FAILED), paid_at")
        ]),
        ("Domain 8: Operations, Notifications & Support", [
            ("campaign_events", "id (PK), campaign_id (FK), event_type (INSTALLED, LIVE, INSPECTED, REMOVED), description, event_time"),
            ("performance_metrics", "id (PK), campaign_id (FK), space_id (FK), metric_date, impressions_count, interactions_count, occupancy_hours"),
            ("notifications", "id (PK), user_id (FK), title, message, type (INFO, WARNING, SUCCESS, ERROR), is_read (BOOL), created_at"),
            ("complaints", "id (PK), campaign_id (FK), user_id (FK), subject, description, priority (LOW, MEDIUM, HIGH), status (OPEN, IN_PROGRESS, RESOLVED), resolved_at")
        ])
    ]

    for domain_name, tables in db_groups:
        add_h2(domain_name)
        for tbl_name, tbl_cols in tables:
            add_bullet(f"Table `{tbl_name}`: ", tbl_cols)

    doc.add_page_break()

    # =========================================================================
    # SECTION 10: FACULTY DEFENSE & VIVA VOCE MASTERCLASS (25+ Q&A)
    # =========================================================================
    add_h1("10. Faculty Defense & Viva Voce Masterclass (25+ Q&A)")
    add_p("This section prepares you for high-intensity faculty questioning. Memorize and understand these technical responses to explain the engineering rigor behind your project.")

    viva_qa = [
        ("Q1: Why did you choose Flask over Django for this project?",
         "Flask is an explicit micro-framework that gives full control over application architecture. Django includes a built-in ORM, admin panel, and template engine that encourages tight coupling and monolithic design. Because we built a modern Decoupled REST API with a React 19 SPA, Flask allowed us to implement the Application Factory Pattern, select PostgreSQL with SQLAlchemy 3, use Marshmallow for precise schema validation, and structure 14 modular Blueprints without carrying Django's monolithic overhead."),

        ("Q2: How does your system guarantee zero double-booking?",
         "We enforce zero double-booking at two complementary levels: (1) Mathematical Interval Collision Detection: When a date range [S_req, E_req] is submitted, we query the space_availability table for any existing booked slot where (S_exist <= E_req AND E_exist >= S_req). (2) PostgreSQL Pessimistic Row-Level Locking: We execute `AdvertisingSpace.query.filter_by(id=space_id).with_for_update().first()`. This locks the inventory row in PostgreSQL. If two users attempt to book the exact same space simultaneously, the second transaction is queued at the database level until the first completes, completely preventing race conditions."),

        ("Q3: What is the Application Factory Pattern, and why is it used?",
         "The Application Factory Pattern (`create_app()`) is a design pattern where the Flask application instance is constructed inside a function rather than as a global variable. It solves circular dependency issues between models, blueprints, and extensions. It also allows spawning multiple instances of the application with different configurations (e.g., DevelopmentConfig vs TestingConfig with an in-memory test database) without process restart."),

        ("Q4: Why do we have extensions.py separate from app/__init__.py?",
         "If extensions like `db = SQLAlchemy()` were instantiated inside `app/__init__.py`, any model importing `db` would also import `app`, creating an immediate circular import loop when `__init__.py` imports the models. By declaring unattached extensions inside `extensions.py` and calling `db.init_app(app)` inside `create_app()`, we break the cycle cleanly."),

        ("Q5: What is the difference between Models and Schemas?",
         "Models (SQLAlchemy) define the database schema—how data is physically stored in PostgreSQL tables with primary/foreign keys and relationships. Schemas (Marshmallow) define the HTTP communication contract. Schemas validate client input data before touching the database, deserialize strings into native types, prevent mass assignment attacks, and serialize database entities into sanitized JSON responses."),

        ("Q6: Why did you use JWT instead of traditional server session cookies?",
         "Traditional session cookies require server-side state (session files or Redis stores) and suffer from Cross-Site Request Forgery (CSRF) vulnerabilities. JSON Web Tokens (JWT) are stateless, cryptographically signed tokens containing user claims (id, role). The frontend stores the token and passes it in the `Authorization: Bearer` header. This allows the backend to remain completely stateless, horizontally scalable, and natively compatible with mobile apps and external APIs."),

        ("Q7: How does the frontend Axios interceptor handle token expiration?",
         "In `services/apiClient.js`, we configure a response interceptor that listens for HTTP 401 Unauthorized errors. When a 401 error occurs, the interceptor automatically flushes the expired `access_token` and `user` profile from `localStorage` and executes `window.location.href = '/login'`. This ensures the user is securely redirected to re-authenticate without corrupting application state."),

        ("Q8: How do you enforce Role-Based Access Control (RBAC) on the backend?",
         "We implement custom Python decorators like `@roles_required('Administrator', 'Space Manager')` on route functions. When a request hits the endpoint, the decorator extracts the user ID from the verified JWT token, queries the user's assigned Role in PostgreSQL, and verifies whether the user's role matches the allowed list. If unauthorized, it halts execution immediately and returns a standardized HTTP 403 Forbidden JSON error."),

        ("Q9: What happens if an advertiser's creative artwork is rejected?",
         "The Creative Reviewer selects REJECTED status and submits a mandatory rejection reason (e.g., 'Dimensions must be 1920x1080; text violates contrast guidelines'). The backend logs the rejection, dispatches an automated notification to the advertiser, and prevents campaign activation. The advertiser can view the feedback on their dashboard and submit a versioned re-upload (version increments from 1 -> 2)."),

        ("Q10: How does the system handle split or partial payments?",
         "The `invoices` table tracks `subtotal`, `tax`, `total_amount`, `amount_paid`, and `balance_due`. When an advertiser pays an installment (e.g., 50% deposit), a record is created in the `payments` table. In the same database transaction, `InvoiceService.reconcile_status(invoice)` sums all completed payments: if `0 < amount_paid < total_amount`, status transitions to `PARTIALLY_PAID`; once `amount_paid >= total_amount`, status flips to `PAID`."),

        ("Q11: Why is Alembic / Flask-Migrate necessary instead of db.create_all()?",
         "`db.create_all()` only creates tables that do not exist; it cannot alter existing tables, rename columns, or drop obsolete indexes without wiping existing production data. Flask-Migrate (Alembic) provides version-controlled database migrations. Every schema modification is recorded as a versioned migration script that can be applied (`flask db upgrade`) or reversed (`flask db downgrade`) safely across development, testing, and production databases."),

        ("Q12: What prevents an advertiser from viewing another advertiser's campaigns?",
         "Multi-tenant data isolation is enforced at the service layer. In `CampaignService.get_all()` and `BookingService.get_all()`, if the authenticated user's role is 'Advertiser', the query automatically appends a filter: `.filter(Campaign.user_id == current_user.id)`. Advertisers cannot access or enumerate other clients' campaigns, bookings, or invoices."),

        ("Q13: What happens when a campaign reaches its end_date?",
         "The system includes automated lifecycle tracking. When a campaign's `end_date` elapses, its status transitions from `ACTIVE` to `COMPLETED`, releasing linked inventory blocks so spaces become available for future bookings. The advertiser's dashboard moves the campaign to historical reporting."),

        ("Q14: How are database transactions protected against unexpected crashes?",
         "Our centralized error handler in `app/errors/handlers.py` catches all `SQLAlchemyError` and unhandled exceptions. In the event of any runtime failure, it immediately invokes `db.session.rollback()` before logging the traceback. This guarantees that partial or corrupted data is never committed to PostgreSQL."),

        ("Q15: What is the purpose of the audit_logs table?",
         "The `audit_logs` table provides an immutable forensic record of all write actions across the system. It logs who performed the action (`user_id`), the action type (`CREATE`, `UPDATE`, `DELETE`, `LOGIN`), the entity type (`Booking`, `Invoice`, `AdvertisingSpace`), snapshots of `old_values` and `new_values` in JSON format, client IP address, and timestamp. Only Administrators can inspect this ledger."),

        ("Q16: How do you calculate pricing for a booking?",
         "Pricing is calculated at the service layer: `duration_days = (end_date - start_date).days + 1`. The base rate from the space's active rate card is multiplied by duration: `total_price = Decimal(space.base_rate) * Decimal(duration_days)`. When formal quotations are generated, sales tax (e.g., 16%) and approved percentage discounts are factored in."),

        ("Q17: Why did you use Three.js in the frontend?",
         "We integrated Three.js and React Three Fiber to provide a real-time 3D billboard preview module. It allows prospective advertisers to upload their 2D digital artwork and visualize it projected onto a photorealistic 3D billboard structure in an urban environment before booking, providing a compelling competitive edge during client pitches."),

        ("Q18: What is the purpose of Flasgger / Swagger in your backend?",
         "Flasgger parses YAML docstrings on every Flask route and automatically generates an interactive OpenAPI 3.0 specification accessible at `http://localhost:5000/apidocs/`. It allows developers and external partners to test every API endpoint, payload, and response code directly in the browser with Bearer token authentication."),

        ("Q19: How are files stored securely in the uploads/ directory?",
         "Uploaded files are not stored with raw client filenames, which could contain path traversal characters (e.g., `../../etc/passwd`). We use `werkzeug.utils.secure_filename()` and append a unique UUID reference string. Stored file paths are recorded in PostgreSQL while the actual binary file resides in the server filesystem."),

        ("Q20: What is the purpose of CORS in Flask, and how is it configured?",
         "Cross-Origin Resource Sharing (CORS) is a browser security mechanism that blocks web pages from making AJAX requests to a different domain or port. Because our React frontend runs on port 5173 and our Flask backend runs on port 5000, we initialize `flask_cors.CORS(app)` with explicit whitelisted origins (`http://localhost:5173`, `http://localhost:3000`) and credentials support."),

        ("Q21: Can an Administrator bypass role checks?",
         "Yes. In `app/common/decorators.py`, the `@roles_required` decorator includes a master bypass check: `if user_role == 'Administrator': return view_function(*args, **kwargs)`. This ensures that platform administrators retain universal oversight for maintenance and customer support."),

        ("Q22: What is the difference between soft delete and hard delete in your system?",
         "Critical business entities like `AdvertisingSpace` and `Campaign` use status flags (`is_active=False` or `status='CANCELLED'`) rather than hard SQL `DELETE` operations. This preserves referential integrity with historical invoices, payments, and audit records."),

        ("Q23: How do you handle password security?",
         "Passwords are never stored in plaintext. We utilize `werkzeug.security.generate_password_hash()` using the salted PBKDF2 with SHA-256 algorithm. Password authentication uses `check_password_hash()`, which is resistant to rainbow table and brute force attacks."),

        ("Q24: How does your frontend state management work?",
         "We combine React Context API for global session identity (`AuthContext`) and UI theming (`ThemeContext`) with feature-isolated local component state and Axios service clients for operational data. This prevents unnecessary global re-renders while keeping the application predictable and modular."),

        ("Q25: What is the demonstration flow for faculty?",
         "Our demonstration follows a complete commercial campaign journey: (1) Log in as Space Manager to showcase inventory and availability calendar; (2) Log in as Advertiser to browse spaces, attempt a double-booking to demonstrate the 409 conflict guardrail, create a campaign, and upload a creative; (3) Log in as Creative Reviewer to inspect and approve the artwork; (4) Log in as Finance Officer to issue the invoice and record bank payment; (5) Show the campaign transitioning to ACTIVE status with live KPI charts; (6) Log in as Administrator to inspect the forensic audit log of the entire sequence.")
    ]

    for q, a in viva_qa:
        add_h3(q)
        add_p(a, space_after=6)

    # =========================================================================
    # SECTION 11: PRESENTATION CHECKLIST FOR FACULTY DEFENSE
    # =========================================================================
    doc.add_page_break()
    add_h1("11. Final Demonstration & Presentation Checklist")
    add_p("Follow this step-by-step checklist to execute a flawless project presentation in front of your faculty:")

    checklist_items = [
        ("Step 1: System Boot & Health Verification",
         "1. Open Terminal 1: Navigate to backend/, activate venv, and run `python run.py`. Backend starts on http://localhost:5000.\n"
         "2. Verify health endpoint in browser: http://localhost:5000/api/health -> Displays status: healthy, database: connected.\n"
         "3. Open Terminal 2: Navigate to frontend/ and run `npm run dev`. Vite client starts on http://localhost:5173."),
        
        ("Step 2: Inventory Discovery & Guardrail 1 Demonstration",
         "1. Log in as Advertiser (e.g., advertiser@test.com / Password123).\n"
         "2. Navigate to Spaces Catalog (`/spaces`). Show search filters by Category (Billboard, LED) and Location (City).\n"
         "3. Open Availability Calendar (`/availability`). Show free dates versus red blocked dates.\n"
         "4. Deliberate Failure Demonstration: Attempt to book an already occupied date range. Show the faculty the explicit error banner: 'Conflict detected: Space is already booked from [Date] to [Date]'. Explain that PostgreSQL row-level locking prevented the conflict."),

        ("Step 3: Campaign Creation & Creative Quality Gate",
         "1. Create a new campaign 'Summer Mega Launch 2026' with budget allocation.\n"
         "2. Select an open advertising space and submit the booking.\n"
         "3. Navigate to Creatives (`/creatives`) and upload billboard artwork banner.\n"
         "4. Show campaign status is currently PENDING_APPROVAL. Explain Guardrail 2: the campaign is legally blocked from broadcasting until certified review."),

        ("Step 4: Creative Reviewer Approval",
         "1. Open a new Incognito browser window (or use demo role switch) and log in as Creative Reviewer.\n"
         "2. Navigate to Review Queue (`/creatives`). Show pending asset thumbnail, dimensions, and file format.\n"
         "3. Click 'Approve Artwork'. Status updates to APPROVED with reviewer timestamp."),

        ("Step 5: Financial Invoicing & Payment Reconciliation",
         "1. Log in as Finance Officer.\n"
         "2. Navigate to Invoices (`/payments`). Locate the campaign invoice showing subtotal, tax calculation, and total.\n"
         "3. Record payment: enter transaction reference 'HBL-BANK-90821' and amount.\n"
         "4. Show atomic status reconciliation: invoice flips to PAID, balance due becomes 0.00.\n"
         "5. Show that the campaign now automatically transitions to ACTIVE status!"),

        ("Step 6: Master Governance & Forensic Audit Trail",
         "1. Log in as Administrator.\n"
         "2. Navigate to Audit Logs (`/audit`).\n"
         "3. Demonstrate the complete chronological audit trail: User login -> Booking creation -> Creative approval -> Payment settlement -> Status update.\n"
         "4. Show Swagger API Documentation at http://localhost:5000/apidocs/ to prove backend enterprise standards.")
    ]

    for step_title, step_body in checklist_items:
        add_h2(step_title)
        add_p(step_body, space_after=8)

    # Save document
    output_filename = "Advertisement_Space_Booking_Comprehensive_Study_Guide.docx"
    doc.save(output_filename)
    print(f"Document successfully created: {output_filename}")

    # Also save a copy to the root workspace for convenience
    dest_path = os.path.join(r"c:\Users\zohai\OneDrive\Desktop\adbooking-system", output_filename)
    dest_path_advertisement = os.path.join(r"c:\Users\zohai\OneDrive\Desktop\Advertisement", output_filename)
    try:
        import shutil
        shutil.copy(output_filename, dest_path)
        shutil.copy(output_filename, dest_path_advertisement)
        print(f"Copied to {dest_path} and {dest_path_advertisement}")
    except Exception as e:
        print(f"Copy note: {e}")

if __name__ == "__main__":
    create_full_project_guide()
