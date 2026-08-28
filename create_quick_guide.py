import os
import shutil
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_quick_study_guide():
    doc = Document()

    # Page Margins: 0.75 inch for a clean, compact 2-3 page layout
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.header_distance = Inches(0.4)
        section.footer_distance = Inches(0.4)

    # Color Palette
    COLOR_NAVY = RGBColor(15, 23, 42)      # Deep Navy for main headings
    COLOR_BLUE = RGBColor(37, 99, 235)     # Bright Blue for subheadings
    COLOR_TEXT = RGBColor(30, 41, 59)      # Charcoal for body text
    COLOR_MUTED = RGBColor(100, 116, 139)  # Slate for captions/subtitles

    HEX_HEADER_BG = "1E293B"   # Table header dark slate
    HEX_ROW_ALT = "F8FAFC"     # Alternate light row
    HEX_CALLOUT_BG = "EFF6FF"  # Soft blue callout background
    HEX_BORDER = "3B82F6"      # Callout left border

    def set_cell_shading(cell, hex_color):
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    def set_cell_margins(cell, top=70, bottom=70, left=100, right=100):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
        tcPr.append(tcMar)

    def add_title(title, subtitle):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(title)
        r.font.name = 'Arial'
        r.font.size = Pt(18)
        r.font.bold = True
        r.font.color.rgb = COLOR_NAVY

        p_sub = doc.add_paragraph()
        p_sub.paragraph_format.space_after = Pt(10)
        r_sub = p_sub.add_run(subtitle)
        r_sub.font.name = 'Calibri'
        r_sub.font.size = Pt(10.5)
        r_sub.font.italic = True
        r_sub.font.color.rgb = COLOR_MUTED

    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(11)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(12.5)
        r.font.bold = True
        r.font.color.rgb = COLOR_NAVY
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(10.5)
        r.font.bold = True
        r.font.color.rgb = COLOR_BLUE
        return p

    def add_p(text, bold_prefix="", space_after=3):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            r_bold = p.add_run(bold_prefix)
            r_bold.font.name = 'Calibri'
            r_bold.font.size = Pt(9.5)
            r_bold.font.bold = True
            r_bold.font.color.rgb = COLOR_NAVY
        r = p.add_run(text)
        r.font.name = 'Calibri'
        r.font.size = Pt(9.5)
        r.font.color.rgb = COLOR_TEXT
        return p

    def add_bullet(bold_label, text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.12
        r1 = p.add_run(bold_label)
        r1.font.name = 'Calibri'
        r1.font.size = Pt(9.5)
        r1.font.bold = True
        r1.font.color.rgb = COLOR_NAVY
        r2 = p.add_run(text)
        r2.font.name = 'Calibri'
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = COLOR_TEXT
        return p

    def add_callout(title, body):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        cell = table.cell(0, 0)
        cell.width = Inches(7.0)
        set_cell_shading(cell, HEX_CALLOUT_BG)
        set_cell_margins(cell, top=70, bottom=70, left=120, right=120)

        tcPr = cell._tc.get_or_add_tcPr()
        borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="0" w:color="{HEX_BORDER}"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
        tcPr.append(borders)

        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(1)
        r_t = p.add_run(f"👉 {title}: ")
        r_t.font.name = 'Arial'
        r_t.font.size = Pt(9.5)
        r_t.font.bold = True
        r_t.font.color.rgb = COLOR_BLUE

        r_b = p.add_run(body)
        r_b.font.name = 'Calibri'
        r_b.font.size = Pt(9)
        r_b.font.color.rgb = COLOR_TEXT
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    def add_table(headers, col_widths, rows_data):
        table = doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        # Headers
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].width = col_widths[i]
            hdr_cells[i].text = h
            set_cell_shading(hdr_cells[i], HEX_HEADER_BG)
            set_cell_margins(hdr_cells[i], top=70, bottom=70, left=90, right=90)
            p = hdr_cells[i].paragraphs[0]
            for r in p.runs:
                r.font.name = 'Arial'
                r.font.size = Pt(8.5)
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)

        # Rows
        for r_idx, row in enumerate(rows_data):
            row_cells = table.add_row().cells
            bg = HEX_ROW_ALT if r_idx % 2 == 1 else "FFFFFF"
            for c_idx, val in enumerate(row):
                row_cells[c_idx].width = col_widths[c_idx]
                row_cells[c_idx].text = str(val)
                set_cell_shading(row_cells[c_idx], bg)
                set_cell_margins(row_cells[c_idx], top=50, bottom=50, left=80, right=80)
                p = row_cells[c_idx].paragraphs[0]
                for r in p.runs:
                    r.font.name = 'Calibri'
                    r.font.size = Pt(8.5)
                    r.font.color.rgb = COLOR_TEXT

        tblPr = table._tbl.tblPr
        borders = parse_xml(f'<w:tblBorders {nsdecls("w")}><w:top w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/><w:insideH w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/><w:insideV w:val="none"/><w:left w:val="none"/><w:right w:val="none"/></w:tblBorders>')
        tblPr.append(borders)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # -------------------------------------------------------------
    # DOCUMENT CONTENT
    # -------------------------------------------------------------
    add_title(
        "Project Architecture, Code Files & Libraries Guide",
        "Quick 3-Page Faculty Study Guide — Simple English Explanation of How the System Works"
    )

    # SECTION 1
    add_h1("1. How the System is Built (The Architecture in Plain Words)")
    add_p("This project is built using a modern 3-Tier Architecture. That simply means the project is split into three independent parts that talk to each other over the network:")
    add_bullet("1. The Frontend (Client Side): ", "Built with React 19 and Vite. This is the visual part that runs inside the user's browser. It shows the buttons, forms, tables, charts, and pages. It does NOT talk directly to the database. Instead, whenever a user clicks something (like 'Book Ad Space'), React sends an HTTP request (in JSON format) to the backend.")
    add_bullet("2. The Backend (Server Side): ", "Built with Python and Flask. This is the 'brain' of the project. It listens on port 5000, checks who the user is (using a secure JWT token), verifies business rules (like making sure a billboard is not already booked), and processes data.")
    add_bullet("3. The Database (Storage Side): ", "We use PostgreSQL. It permanently stores all our tables (users, spaces, bookings, campaigns, invoices, audit logs). The Flask backend talks to PostgreSQL using SQLAlchemy, an Object-Relational Mapper (ORM) that translates Python code into database tables.")

    add_callout(
        "How to explain this to faculty",
        "'Sir, our application is completely decoupled. The React frontend handles user interaction, the Flask REST API handles business logic and security, and PostgreSQL stores the data. They communicate purely using JSON requests authenticated with Bearer JWT tokens.'"
    )

    # SECTION 2
    add_h1("2. What Every File & Folder is For (Folder Breakdown)")
    add_p("Here is the simple map of what every key file and folder does in both the backend and frontend:")

    add_h2("A. Backend Files & Folders (backend/)")
    backend_file_data = [
        ("run.py", "Root file for development. Runs `python run.py`. Starts the Flask server with live reload on port 5000."),
        ("wsgi.py", "Production root file. Exposes the app for production servers like Gunicorn."),
        ("config.py", "Holds settings: database URL, JWT secret keys, upload folder paths, and CORS allowed origins."),
        ("requirements.txt", "The list of all Python libraries and packages required to run the backend."),
        (".env", "Stores private passwords and secrets (like database credentials) so they are not hardcoded in code."),
        ("app/__init__.py", "The Application Factory. Defines `create_app()`, connects the database, and registers all 14 blueprints."),
        ("app/extensions.py", "Creates unattached instances of `db = SQLAlchemy()`, `jwt = JWTManager()`, and `migrate = Migrate()` to avoid circular imports."),
        ("app/models/", "Folder containing 15 database model files (user, space, booking, campaign, invoice, etc.) defining PostgreSQL tables."),
        ("app/services/", "Folder containing the pure business logic (BookingService, AvailabilityService, PaymentService). Keeps routes clean."),
        ("app/schemas.py", "Marshmallow schemas that check and validate user input data before it touches the database."),
        ("app/common/decorators.py", "Holds `@roles_required()` to restrict endpoints based on the user's role (e.g. only Admins can manage users)."),
        ("app/errors/handlers.py", "Catches all errors (validation, 404, database crash) and rolls back transactions safely."),
        ("app/blueprints/ (14 folders)", "Modular route folders: auth, spaces, availability, campaigns, bookings, creatives, invoices, payments, dashboard, etc."),
        ("uploads/", "Folder on the server disk where uploaded creative artwork (images, videos, PDFs) is saved.")
    ]
    add_table(
        ["File / Folder Name", "What It Does in Simple English"],
        [Inches(2.2), Inches(4.8)],
        backend_file_data
    )

    add_h2("B. Frontend Files & Folders (frontend/)")
    frontend_file_data = [
        ("package.json", "List of all npm packages and libraries installed for the React frontend."),
        ("vite.config.js", "Configuration for Vite, which builds and runs the React app ultra-fast."),
        ("index.html", "The single HTML page containing `<div id='root'></div>` where React renders."),
        ("src/main.jsx", "The JavaScript entry point. Mounts the React app into the HTML file and loads Bootstrap CSS."),
        ("src/App.jsx", "The master router. Defines all URLs (`/login`, `/dashboard`, `/spaces`, etc.) and wraps them in protected routes."),
        ("src/context/AuthContext.jsx", "Stores the logged-in user info, JWT token, and role helper functions (`isAdmin`, `isAdvertiser`)."),
        ("src/services/apiClient.js", "Configured Axios HTTP client. Automatically attaches the JWT token to every request and redirects to login on 401."),
        ("src/layouts/AppShell.jsx", "The main dashboard layout with the dynamic sidebar, top navbar, theme toggle, and notification bell."),
        ("src/features/ (19 folders)", "Self-contained feature folders (spaces, campaigns, bookings, creatives, payments, etc.) with their own pages.")
    ]
    add_table(
        ["File / Folder Name", "What It Does in Simple English"],
        [Inches(2.2), Inches(4.8)],
        frontend_file_data
    )

    # PAGE BREAK FOR LIBRARIES & CHEAT SHEET
    doc.add_page_break()

    # SECTION 3
    add_h1("3. Libraries Used & Why We Chose Them")
    add_p("Faculty always ask: 'What libraries did you use, and why did you pick this library instead of another one?' Here are the clear, simple answers:")

    add_h2("A. Backend Python Libraries")
    be_lib_data = [
        ("Flask", "Core Web Framework", "Why used: It is a lightweight micro-framework. Unlike Django, which forces its own rigid structure, Flask lets us build a fast, clean REST API with modular blueprints."),
        ("Flask-SQLAlchemy", "Database ORM", "Why used: Allows us to write clean Python code (`User.query.filter_by(...)`) instead of writing messy, dangerous raw SQL queries. Prevents SQL injection."),
        ("Flask-Migrate / Alembic", "Database Migrations", "Why used: It acts like 'Git for database tables'. When we add a new column or table, Alembic upgrades the database safely without deleting existing data."),
        ("Flask-JWT-Extended", "Authentication & Tokens", "Why used: Generates secure JSON Web Tokens. Keeps our API completely stateless, meaning the server doesn't need to store session files in memory."),
        ("Marshmallow", "Data Validation", "Why used: Checks user input before saving to database (e.g. valid email, strong password, positive price). Stops malicious or bad data instantly."),
        ("Flask-CORS", "Cross-Origin Security", "Why used: Browsers block React (port 5173) from talking to Flask (port 5000) for security. Flask-CORS tells the browser to safely allow this communication."),
        ("psycopg (or psycopg2)", "PostgreSQL Driver", "Why used: The official high-speed database driver that connects Python directly to PostgreSQL over TCP port 5432."),
        ("Flasgger (Swagger)", "API Documentation", "Why used: Automatically generates an interactive documentation website at `/apidocs/` where anyone can test the API directly from the browser."),
        ("Werkzeug", "Security & Utilities", "Why used: Flask's built-in utility engine. Used for secure password hashing (`generate_password_hash`) and cleaning upload file names.")
    ]
    add_table(
        ["Library Name", "What It Is", "Why We Chose It (Faculty Answer)"],
        [Inches(1.5), Inches(1.5), Inches(4.0)],
        be_lib_data
    )

    add_h2("B. Frontend React Libraries")
    fe_lib_data = [
        ("React 19", "UI Component Library", "Why used: The industry standard for building fast, component-based user interfaces. Re-renders only the parts of the screen that change, making it very smooth."),
        ("Vite", "Build Tool & Dev Server", "Why used: Replaces older, slow tools like Create-React-App. Starts in less than a second and gives instant Hot Module Reloading (HMR) during coding."),
        ("React Router v7", "Client-Side Routing", "Why used: Allows navigation between pages (Dashboard, Bookings, Creatives) instantly without ever reloading the browser window (Single Page App)."),
        ("Axios", "HTTP Network Client", "Why used: Much better than browser `fetch()`. Supports request interceptors to automatically attach JWT tokens and response interceptors to catch expired sessions."),
        ("Bootstrap 5 & React-Bootstrap", "CSS & UI Styling", "Why used: Provides polished, responsive UI grids, buttons, forms, tables, and modals out of the box so the app looks professional on all screen sizes."),
        ("Chart.js & react-chartjs-2", "Data Visualization", "Why used: Renders interactive charts (revenue graphs, space occupancy rates) on the dashboard so faculty can see visual analytics."),
        ("Three.js & React Three Fiber", "3D WebGL Graphics", "Why used: Renders a real-time 3D billboard on screen! Advertisers can preview their 2D ad image mapped onto a realistic 3D billboard structure."),
        ("Lucide React", "Icon Pack", "Why used: Lightweight, modern vector icons used across the sidebar, buttons, and status badges."),
        ("React Hook Form", "Form Management", "Why used: Handles complex forms (campaign creation, booking inputs) efficiently without causing laggy screen re-renders.")
    ]
    add_table(
        ["Library Name", "What It Is", "Why We Chose It (Faculty Answer)"],
        [Inches(1.5), Inches(1.5), Inches(4.0)],
        fe_lib_data
    )

    # SECTION 4
    add_h1("4. The 5 Quick Golden Rules for Your Faculty Defense")
    add_p("Keep these 5 points in your head when talking to your faculty. They cover 90% of the questions they will ask:")
    add_bullet("1. Why Flask and not Django? ", "'Sir, Django is a heavy monolith designed for full-stack template apps. Since we built a decoupled React SPA, Flask gave us the exact lightweight flexibility we needed for a modular REST API.'")
    add_bullet("2. How do you prevent double-booking? ", "'Sir, two ways: first, we check date intervals mathematically (Existing_Start <= New_End AND Existing_End >= New_Start). Second, we use PostgreSQL row-level locking via SQLAlchemy's `.with_for_update()` so two users clicking at the exact same millisecond cannot collide.'")
    add_bullet("3. What is the difference between Models and Schemas? ", "'Sir, Models define the PostgreSQL database tables. Schemas (Marshmallow) validate the user's incoming JSON data and strip dangerous fields before anything touches the database.'")
    add_bullet("4. Why do we have app/__init__.py and extensions.py? ", "'Sir, `app/__init__.py` uses the Application Factory Pattern (`create_app()`) to create the app dynamically. `extensions.py` holds unattached database objects to completely avoid circular import errors.'")
    add_bullet("5. How does the frontend stay secure? ", "'Sir, we use JWT tokens. Our Axios client automatically adds `Authorization: Bearer <token>` to every request. If the token expires, an Axios interceptor catches the 401 error and safely redirects the user back to the login page.'")

    output_filename = "Project_Quick_Study_Guide_Architecture_and_Libraries.docx"
    doc.save(output_filename)
    print(f"Quick Guide created: {output_filename}")

    # Copy to both workspace directories
    path1 = os.path.join(r"c:\Users\zohai\OneDrive\Desktop\adbooking-system", output_filename)
    path2 = os.path.join(r"c:\Users\zohai\OneDrive\Desktop\Advertisement", output_filename)
    shutil.copy2(output_filename, path1)
    shutil.copy2(output_filename, path2)
    print(f"Copied to {path1} and {path2}")

if __name__ == "__main__":
    create_quick_study_guide()
